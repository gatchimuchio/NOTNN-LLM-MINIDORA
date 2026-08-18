"""ローカル文書の安全な抽出。外部URLは取得しない。"""
from __future__ import annotations

import csv
import hashlib
import json
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from .runtime import DocumentInput, MiniDoraEngine

_TEXT = {".txt", ".md", ".markdown", ".rst", ".log"}
_SUPPORTED = _TEXT | {".csv", ".json", ".jsonl", ".html", ".htm", ".pdf", ".docx"}


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    title: str
    body: str
    source_uri: str
    metadata: dict[str, Any]


class _HTML(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.suppressed = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.casefold() in {"script", "style", "noscript"}:
            self.suppressed += 1
        elif tag.casefold() in {"p", "br", "div", "li", "h1", "h2", "h3", "tr"}:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in {"script", "style", "noscript"} and self.suppressed:
            self.suppressed -= 1
        elif tag.casefold() in {"p", "div", "li", "h1", "h2", "h3", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.suppressed and data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(line for line in (" ".join(row.split()) for row in "".join(self.parts).splitlines()) if line)


def discover(path: Path, *, recursive: bool) -> list[Path]:
    expanded = path.expanduser()
    if expanded.is_symlink():
        raise ValueError(f"symlinkは投入できません: {expanded}")
    root = expanded.resolve()
    if root.is_file():
        return [root]
    if not root.is_dir():
        raise ValueError(f"pathが存在しません: {root}")
    iterator = root.rglob("*") if recursive else root.glob("*")
    return sorted(item for item in iterator if item.is_file() and not item.is_symlink() and item.suffix.casefold() in _SUPPORTED)


def extract(path: Path, *, max_bytes: int = 20 * 1024 * 1024) -> ExtractedDocument:
    if path.expanduser().is_symlink():
        raise ValueError(f"symlinkは投入できません: {path}")
    resolved = path.expanduser().resolve()
    if not resolved.is_file():
        raise ValueError(f"通常fileだけを投入できます: {resolved}")
    size = resolved.stat().st_size
    if size <= 0 or size > max_bytes:
        raise ValueError(f"文書sizeが境界外です: {size}")
    suffix = resolved.suffix.casefold()
    if suffix not in _SUPPORTED:
        raise ValueError(f"未対応形式です: {suffix}")
    raw = resolved.read_bytes()
    body = _extract(resolved, raw, suffix).replace("\x00", "").strip()
    if not body:
        raise ValueError(f"本文を抽出できません: {resolved}")
    return ExtractedDocument(
        resolved.stem,
        body,
        resolved.as_uri(),
        {"format": suffix.removeprefix("."), "filename": resolved.name, "size_bytes": size, "sha256": hashlib.sha256(raw).hexdigest()},
    )


def ingest(engine: MiniDoraEngine, path: Path, *, recursive: bool = False, max_bytes: int = 20 * 1024 * 1024) -> list[str]:
    ids: list[str] = []
    for candidate in discover(path, recursive=recursive):
        item = extract(candidate, max_bytes=max_bytes)
        ids.append(engine.add_document(DocumentInput(item.title, item.body, item.source_uri, item.metadata)))
    if not ids:
        raise ValueError("投入可能な文書fileがありません")
    return ids


def _decode(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文字encodingを判定できません")


def _json_text(value: Any, prefix: str = "") -> str:
    if isinstance(value, dict):
        return "\n".join(_json_text(value[key], f"{prefix}.{key}" if prefix else str(key)) for key in sorted(value, key=str))
    if isinstance(value, list):
        return "\n".join(_json_text(item, f"{prefix}[{index}]") for index, item in enumerate(value))
    return f"{prefix}: {value}" if prefix else str(value)


def _extract(path: Path, raw: bytes, suffix: str) -> str:
    if suffix in _TEXT:
        return _decode(raw)
    if suffix == ".csv":
        return "\n".join("\t".join(cell.strip() for cell in row) for row in csv.reader(_decode(raw).splitlines()))
    if suffix == ".json":
        return _json_text(json.loads(_decode(raw)))
    if suffix == ".jsonl":
        return "\n".join(_json_text(json.loads(line)) for line in _decode(raw).splitlines() if line.strip())
    if suffix in {".html", ".htm"}:
        parser = _HTML(); parser.feed(_decode(raw)); return parser.text()
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ValueError("PDF投入にはdocuments extraが必要です") from exc
        return "\n\n".join((page.extract_text() or "").strip() for page in PdfReader(path).pages)
    if suffix == ".docx":
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ValueError("DOCX投入にはdocuments extraが必要です") from exc
        return "\n".join(paragraph.text.strip() for paragraph in Document(path).paragraphs if paragraph.text.strip())
    raise ValueError(f"未対応形式です: {suffix}")
